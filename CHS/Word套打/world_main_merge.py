import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkinter.scrolledtext import ScrolledText
import openpyxl
import re
import os
import shutil
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class WordMailMergeApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Word套打工具")
        self.root.geometry("900x750")
        
        self.word_template_path = ""
        self.excel_data_path = ""
        self.excel_sheets = []
        self.markers = []
        self.column_names = []
        self.marker_column_mapping = {}
        self.filename_column = ""
        
        self.create_widgets()
    
    def create_widgets(self):
        # Word模板选择
        word_frame = ttk.LabelFrame(self.root, text="Word模板文件", padding=10)
        word_frame.pack(fill=tk.X, padx=10, pady=5)
        
        self.word_path_var = tk.StringVar()
        word_entry = ttk.Entry(word_frame, textvariable=self.word_path_var, state="readonly")
        word_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 10))
        ttk.Button(word_frame, text="选择Word模板", command=self.select_word_template).pack(side=tk.LEFT)
        
        # Excel数据源选择
        excel_frame = ttk.LabelFrame(self.root, text="Excel数据源", padding=10)
        excel_frame.pack(fill=tk.X, padx=10, pady=5)
        
        self.excel_path_var = tk.StringVar()
        excel_entry = ttk.Entry(excel_frame, textvariable=self.excel_path_var, state="readonly")
        excel_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 10))
        ttk.Button(excel_frame, text="选择Excel文件", command=self.select_excel_file).pack(side=tk.LEFT)
        
        ttk.Label(excel_frame, text="选择Sheet:").pack(side=tk.LEFT, padx=(20, 5))
        self.sheet_var = tk.StringVar()
        self.sheet_combobox = ttk.Combobox(excel_frame, textvariable=self.sheet_var, state="readonly", width=25)
        self.sheet_combobox.pack(side=tk.LEFT)
        self.sheet_combobox.bind("<<ComboboxSelected>>", self.on_sheet_selected)
        
        # 标记映射区域
        mapping_frame = ttk.LabelFrame(self.root, text="标记与列映射", padding=10)
        mapping_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        # 创建表格显示标记和对应的列
        columns_frame = ttk.Frame(mapping_frame)
        columns_frame.pack(fill=tk.BOTH, expand=True)
        
        ttk.Label(columns_frame, text="Word标记", font=("Arial", 10, "bold")).grid(row=0, column=0, sticky="w", padx=5, pady=5)
        ttk.Label(columns_frame, text="对应Excel列", font=("Arial", 10, "bold")).grid(row=0, column=1, sticky="w", padx=5, pady=5)
        
        # 标记列表和下拉框的容器
        self.mapping_container = ttk.Frame(columns_frame)
        self.mapping_container.grid(row=1, column=0, columnspan=2, sticky="nsew")
        columns_frame.columnconfigure(0, weight=1)
        columns_frame.rowconfigure(1, weight=1)
        
        # 文件名列选择
        filename_frame = ttk.LabelFrame(self.root, text="文件名列选择", padding=10)
        filename_frame.pack(fill=tk.X, padx=10, pady=5)
        
        ttk.Label(filename_frame, text="选择用于生成文件名的列:").pack(side=tk.LEFT, padx=5)
        self.filename_column_var = tk.StringVar()
        self.filename_combobox = ttk.Combobox(filename_frame, textvariable=self.filename_column_var, state="readonly", width=30)
        self.filename_combobox.pack(side=tk.LEFT, padx=5)
        
        # 执行按钮
        button_frame = ttk.Frame(self.root, padding=10)
        button_frame.pack(fill=tk.X, padx=10, pady=5)
        
        ttk.Button(button_frame, text="执行套打", command=self.execute_mail_merge, padding=10).pack(side=tk.RIGHT)
        
        # 日志区域
        log_frame = ttk.LabelFrame(self.root, text="执行日志", padding=10)
        log_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        self.log_text = ScrolledText(log_frame, height=10, wrap=tk.WORD)
        self.log_text.pack(fill=tk.BOTH, expand=True)
    
    def log(self, message):
        self.log_text.insert(tk.END, message + "\n")
        self.log_text.see(tk.END)
        self.root.update()
    
    def select_word_template(self):
        file_path = filedialog.askopenfilename(
            title="选择Word模板文件",
            filetypes=[("Word文档", "*.doc *.docx")]
        )
        if file_path:
            self.word_template_path = file_path
            self.word_path_var.set(file_path)
            self.log(f"已选择Word模板: {file_path}")
            self.load_word_markers()
    
    def select_excel_file(self):
        file_path = filedialog.askopenfilename(
            title="选择Excel数据文件",
            filetypes=[("Excel文件", "*.xlsx *.xls")]
        )
        if file_path:
            self.excel_data_path = file_path
            self.excel_path_var.set(file_path)
            self.log(f"已选择Excel文件: {file_path}")
            self.load_excel_sheets()
    
    def load_word_markers(self):
        if not self.word_template_path:
            return
        
        try:
            doc = Document(self.word_template_path)
            markers_dict = {}
            
            # 搜索段落中的标记，记录出现顺序
            for para in doc.paragraphs:
                matches = re.finditer(r'【([^】]+)】', para.text)
                for match in matches:
                    marker = match.group(1)
                    if marker not in markers_dict:
                        markers_dict[marker] = len(markers_dict)
            
            # 搜索表格中的标记，记录出现顺序
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        matches = re.finditer(r'【([^】]+)】', cell.text)
                        for match in matches:
                            marker = match.group(1)
                            if marker not in markers_dict:
                                markers_dict[marker] = len(markers_dict)
            
            # 按出现顺序排序标记
            self.markers = sorted(markers_dict.keys(), key=lambda x: markers_dict[x])
            self.log(f"在Word模板中找到 {len(self.markers)} 个标记: {', '.join(self.markers)}")
            
            self.refresh_mapping_ui()
            self.refresh_filename_combobox()
            
        except Exception as e:
            messagebox.showerror("错误", f"读取Word文件失败: {str(e)}")
            self.log(f"错误: {str(e)}")
    
    def load_excel_sheets(self):
        if not self.excel_data_path:
            return
        
        try:
            wb = openpyxl.load_workbook(self.excel_data_path, read_only=True)
            self.excel_sheets = wb.sheetnames
            wb.close()
            
            if self.excel_sheets:
                self.sheet_combobox['values'] = self.excel_sheets
                self.sheet_combobox.set(self.excel_sheets[0])
                self.log(f"Excel文件包含 {len(self.excel_sheets)} 个Sheet: {', '.join(self.excel_sheets)}")
                self.load_excel_columns()
            else:
                messagebox.showwarning("警告", "Excel文件中没有Sheet")
                self.log("警告: Excel文件中没有Sheet")
                
        except Exception as e:
            messagebox.showerror("错误", f"读取Excel文件失败: {str(e)}")
            self.log(f"错误: {str(e)}")
    
    def on_sheet_selected(self, event=None):
        selected_sheet = self.sheet_var.get()
        if selected_sheet:
            self.log(f"已选择Sheet: {selected_sheet}")
            self.load_excel_columns()
    
    def load_excel_columns(self):
        if not self.excel_data_path or not self.sheet_var.get():
            return
        
        try:
            # 使用data_only=True获取公式的计算结果
            wb = openpyxl.load_workbook(self.excel_data_path, read_only=True, data_only=True)
            ws = wb[self.sheet_var.get()]
            
            # 读取第一行作为列名
            self.column_names = []
            for cell in ws[1]:
                if cell.value:
                    self.column_names.append(str(cell.value))
            
            wb.close()
            self.log(f"Excel文件中有 {len(self.column_names)} 列: {', '.join(self.column_names)}")
            
            self.refresh_mapping_ui()
            self.refresh_filename_combobox()
            
        except Exception as e:
            messagebox.showerror("错误", f"读取Excel文件失败: {str(e)}")
            self.log(f"错误: {str(e)}")
    
    def refresh_mapping_ui(self):
        # 清空现有的映射界面
        for widget in self.mapping_container.winfo_children():
            widget.destroy()
        
        self.marker_column_mapping = {}
        
        # 为每个标记创建一行
        for idx, marker in enumerate(self.markers):
            # 标记显示
            marker_label = ttk.Label(self.mapping_container, text=f"【{marker}】")
            marker_label.grid(row=idx, column=0, sticky="w", padx=5, pady=3)
            
            # 下拉框
            column_var = tk.StringVar()
            column_combobox = ttk.Combobox(
                self.mapping_container,
                textvariable=column_var,
                values=["(不指定)"] + self.column_names,
                state="readonly",
                width=40
            )
            column_combobox.set("(不指定)")
            column_combobox.grid(row=idx, column=1, sticky="w", padx=5, pady=3)
            
            # 保存映射关系
            self.marker_column_mapping[marker] = column_var
    
    def refresh_filename_combobox(self):
        if self.column_names:
            self.filename_combobox['values'] = self.column_names
            if self.filename_column_var.get() not in self.column_names:
                self.filename_combobox.set("")
        else:
            self.filename_combobox['values'] = []
            self.filename_column_var.set("")
    
    def replace_text_keep_format(self, doc, old_text, new_text):
        """替换文本同时保持格式（包括下划线等）"""
        new_text = str(new_text)
        
        # 处理段落中的替换
        for para in doc.paragraphs:
            if old_text in para.text:
                self._replace_in_paragraph(para, old_text, new_text)
        
        # 处理表格中的替换
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if old_text in para.text:
                            self._replace_in_paragraph(para, old_text, new_text)
        
        return doc
    
    def _replace_in_paragraph(self, para, old_text, new_text):
        """在段落中替换文本，保持原有格式"""
        runs = para.runs
        if not runs:
            return
        
        # 构建完整的文本和对应的格式信息
        full_text = ''.join(run.text for run in runs)
        
        if old_text not in full_text:
            return
        
        # 找到所有需要替换的位置
        old_pos = 0
        replacement_positions = []
        while True:
            pos = full_text.find(old_text, old_pos)
            if pos == -1:
                break
            replacement_positions.append((pos, pos + len(old_text)))
            old_pos = pos + 1
        
        if not replacement_positions:
            return
        
        # 计算每个字符属于哪个run
        char_to_run = []
        current_pos = 0
        for run_idx, run in enumerate(runs):
            run_text = run.text
            for _ in run_text:
                char_to_run.append(run_idx)
                current_pos += 1
        
        # 替换文本
        offset = 0
        for start_pos, end_pos in replacement_positions:
            real_start = start_pos + offset
            real_end = end_pos + offset
            
            # 找到标记所在的run
            start_run_idx = char_to_run[real_start] if real_start < len(char_to_run) else -1
            
            if start_run_idx != -1:
                # 在该run中直接替换
                runs[start_run_idx].text = runs[start_run_idx].text.replace(old_text, new_text)
            
            # 更新偏移量（因为替换后文本长度可能变化）
            offset += len(new_text) - len(old_text)
        
        # 清空内容并重新构建段落
        new_runs = []
        new_runs_text = []
        
        for run in runs:
            if run.text:
                new_runs.append(run)
                new_runs_text.append(run.text)
        
        # 重建段落内容
        para.clear()
        for run, text in zip(new_runs, new_runs_text):
            new_run = para.add_run(text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.name = run.font.name
            new_run.font.size = run.font.size
            new_run.font.color.rgb = run.font.color.rgb
    
    def execute_mail_merge(self):
        # 验证输入
        if not self.word_template_path:
            messagebox.showwarning("警告", "请先选择Word模板文件")
            return
        
        if not self.excel_data_path:
            messagebox.showwarning("警告", "请先选择Excel数据文件")
            return
        
        # 验证文件名列
        if not self.filename_column_var.get():
            messagebox.showwarning("警告", "必须指定文件名列")
            return
        
        # 选择输出文件夹
        output_folder = filedialog.askdirectory(title="选择输出文件夹")
        if not output_folder:
            return
        
        self.log("=" * 50)
        self.log("开始执行套打...")
        
        try:
            # 读取Excel数据，使用data_only=True获取公式的计算结果
            wb = openpyxl.load_workbook(self.excel_data_path, read_only=True, data_only=True)
            ws = wb[self.sheet_var.get()]
            
            # 获取列名和列索引映射
            header_row = []
            for cell in ws[1]:
                if cell.value:
                    header_row.append(str(cell.value))
            
            # 创建列名到列索引的映射（从1开始）
            column_index_map = {name: idx + 1 for idx, name in enumerate(header_row)}
            
            # 获取文件名列索引（必须指定）
            filename_col_index = column_index_map.get(self.filename_column_var.get())
            
            # 获取标记到列的映射
            mapping = {}
            for marker, var in self.marker_column_mapping.items():
                selected_column = var.get()
                if selected_column != "(不指定)" and selected_column in column_index_map:
                    mapping[marker] = column_index_map[selected_column]
            
            self.log(f"标记映射: {mapping}")
            self.log(f"文件名列: {self.filename_column_var.get()}")
            
            # 获取数据行（从第2行开始），使用data_only=True获取公式的计算结果
            data_rows = list(ws.iter_rows(min_row=2, values_only=True))
            wb.close()
            
            if not data_rows:
                messagebox.showwarning("警告", "Excel文件中没有数据行")
                return
            
            self.log(f"共 {len(data_rows)} 条数据待处理")
            
            # 为每一行生成Word文件
            success_count = 0
            for row_idx, row in enumerate(data_rows, start=2):
                try:
                    # 读取数据（values_only=True，row已经是值列表）
                    row_data = {}
                    for col_name, col_idx in column_index_map.items():
                        cell_value = row[col_idx - 1]
                        row_data[col_name] = cell_value if cell_value is not None else ""
                    
                    # 生成文件名（必须指定）
                    filename_value = row[filename_col_index - 1]
                    if filename_value:
                        base_name = str(filename_value)
                    else:
                        base_name = f"output_{row_idx}"
                    
                    # 清理文件名中的非法字符
                    base_name = re.sub(r'[<>:"/\\|?*]', '_', base_name)
                    output_filename = f"{base_name}.docx"
                    output_path = os.path.join(output_folder, output_filename)
                    
                    # 复制模板
                    shutil.copy2(self.word_template_path, output_path)
                    
                    # 替换标记
                    doc = Document(output_path)
                    
                    for marker, col_idx in mapping.items():
                        old_text = f"【{marker}】"
                        new_value = row[col_idx - 1]
                        new_text = str(new_value) if new_value is not None else ""
                        
                        self.replace_text_keep_format(doc, old_text, new_text)
                    
                    # 保存文档
                    doc.save(output_path)
                    
                    self.log(f"成功生成: {output_filename}")
                    success_count += 1
                    
                except Exception as e:
                    self.log(f"处理第{row_idx}行时出错: {str(e)}")
            
            self.log("=" * 50)
            self.log(f"套打完成! 成功生成 {success_count}/{len(data_rows)} 个文件")
            messagebox.showinfo("完成", f"套打完成!\n成功生成 {success_count}/{len(data_rows)} 个文件")
            
        except Exception as e:
            self.log(f"执行过程中发生错误: {str(e)}")
            messagebox.showerror("错误", f"执行失败: {str(e)}")


def main():
    root = tk.Tk()
    app = WordMailMergeApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
